import json
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import sys

sys.stdout.reconfigure(encoding='utf-8')

JSON_PATH = r"C:\Antigravity\Bhagavadgita\_extracted\clean_verses_700.json"
DOCX_PATH = r"C:\Antigravity\Bhagavadgita\Bhagavad_Gita_All_Verses_CLEAN.docx"

CHAPTER_NAMES = {
    1: "ಅರ್ಜುನವಿಷಾದ ಯೋಗ",
    2: "ಸಾಂಖ್ಯ ಯೋಗ",
    3: "ಕರ್ಮ ಯೋಗ",
    4: "ಜ್ಞಾನಕರ್ಮಸನ್ಯಾಸ ಯೋಗ",
    5: "ಕರ್ಮಸನ್ಯಾಸ ಯೋಗ",
    6: "ಆತ್ಮಸಂಯಮ ಯೋಗ",
    7: "ಜ್ಞಾನವಿಜ್ಞಾನ ಯೋಗ",
    8: "ಅಕ್ಷರಪರಬ್ರಹ್ಮ ಯೋಗ",
    9: "ರಾಜವಿದ್ಯಾರಾಜಗುಹ್ಯ ಯೋಗ",
    10: "ವಿಭೂತಿ ಯೋಗ",
    11: "ವಿಶ್ವರೂಪದರ್ಶನ ಯೋಗ",
    12: "ಭಕ್ತಿ ಯೋಗ",
    13: "ಕ್ಷೇತ್ರಕ್ಷೇತ್ರಜ್ಞವಿಭಾಗ ಯೋಗ",
    14: "ಗುಣತ್ರಯವಿಭಾಗ ಯೋಗ",
    15: "ಪುರುಷೋತ್ತಮ ಯೋಗ",
    16: "ದೈವಾಸುರಸಂಪದ್ವಿಭಾಗ ಯೋಗ",
    17: "ಶ್ರದ್ಧಾತ್ರಯವಿಭಾಗ ಯೋಗ",
    18: "ಮೋಕ್ಷಸನ್ಯಾಸ ಯೋಗ"
}

def set_cell_background(cell, hex_color):
    """Sets the background color (shading) of a table cell."""
    shd_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shd_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets internal padding (margins) of a cell in dxa (1/20 of a pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def build_docx():
    print(f"Loading {JSON_PATH}...")
    with open(JSON_PATH, "r", encoding="utf-8") as f:
        data = json.load(f)
    print(f"Loaded {len(data)} verses.")

    doc = Document()
    
    # Configure document margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Setup core fonts
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Segoe UI'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x1F, 0x1A, 0x13) # Warm dark brown ink

    # ---------------- Title Page ----------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("\n\n\n\nಭಗವದ್ಗೀತಾ\n")
    title_run.font.name = 'Tiro Kannada'
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x8B, 0x1E, 0x1A) # Crimson deep red

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("ಸಂಸ್ಕೃತ ಶ್ಲೋಕ ಮತ್ತು ಶ್ರೀ ಬನ್ನಂಜೆ ಗೋವಿಂದಾಚಾರ್ಯರ ಕನ್ನಡ ವಿವರಣೆ\n\n")
    sub_run.font.name = 'Tiro Kannada'
    sub_run.font.size = Pt(14)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(0x4A, 0x3F, 0x30)

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_p.add_run("ಸಂಪೂರ್ಣ ಹದಿನೆಂಟು ಅಧ್ಯಾಯಗಳು\n\n\n\n\n\n")
    meta_run.font.name = 'Segoe UI'
    meta_run.font.size = Pt(12)
    meta_run.font.bold = True
    meta_run.font.color.rgb = RGBColor(0x7A, 0x6C, 0x54)
    
    attribution_p = doc.add_paragraph()
    attribution_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    attr_run = attribution_p.add_run("ಮೂಲ ಕೃತಿ ಆಧಾರ: bhagavadgitakannada.blogspot.com\nವೈಯಕ್ತಿಕ ಅಧ್ಯಯನಕ್ಕಾಗಿ ಮಾತ್ರ")
    attr_run.font.name = 'Segoe UI'
    attr_run.font.size = Pt(10)
    attr_run.font.color.rgb = RGBColor(0x7A, 0x6C, 0x54)

    doc.add_page_break()

    # ---------------- Table of Contents & Status Legend ----------------
    toc_head = doc.add_paragraph()
    toc_run = toc_head.add_run("ಅಧ್ಯಾಯಗಳು ಮತ್ತು ವಿವರಣೆ ಸೂಚಿ (Legend)")
    toc_run.font.name = 'Tiro Kannada'
    toc_run.font.size = Pt(18)
    toc_run.font.bold = True
    toc_run.font.color.rgb = RGBColor(0x8B, 0x1E, 0x1A)
    
    doc.add_paragraph("ಈ ಕೃತಿಯಲ್ಲಿನ ಶ್ಲೋಕಗಳು ಮತ್ತು ವಿವರಣೆಗಳನ್ನು ಕೆಳಗಿನ ಸ್ಥಿತಿಗಳ (Status Flags) ಆಧಾರದ ಮೇಲೆ ವರ್ಗೀಕರಿಸಲಾಗಿದೆ ಮತ್ತು ಅವುಗಳನ್ನು ಸುಲಭವಾಗಿ ಗುರುತಿಸಲು ಪ್ರತ್ಯೇಕ ಬಣ್ಣದ ಕಾರ್ಡ್‌ಗಳಲ್ಲಿ ನೀಡಲಾಗಿದೆ:\n")
    
    # Legend Table
    legend_table = doc.add_table(rows=5, cols=3)
    legend_table.autofit = False
    
    headers = ["ಸ್ಥಿತಿ (Status)", "ಬಣ್ಣ (Color)", "ವಿವರಣೆ (Description)"]
    for i, h in enumerate(headers):
        cell = legend_table.cell(0, i)
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_background(cell, "D9D9D9")
        set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
        
    legend_data = [
        ("clean", "Warm Cream (ಕ್ಲೀನ್)", " printed ಪುಸ್ತಕದೊಂದಿಗೆ ಸಂಪೂರ್ಣವಾಗಿ ತಾಳೆನೋಡಿ ಪರಿಶೀಲಿಸಿದ ಪರಿಶುದ್ಧ ಮುದ್ರಿತ OCR ಪಠ್ಯ."),
        ("screenshot_patch", "Soft Green (ಹಸಿರು)", "ಬಳಕೆದಾರರು ಕಳುಹಿಸಿದ ಮುದ್ರಿತ ಪುಸ್ತಕದ ಪುಟಗಳ ಸ್ಕ್ರೀನ್‌ಶಾಟ್‌ಗಳಿಂದ ನೇರವಾಗಿ ಪ್ರತಿಲಿಖಿಸಿ ದಾಖಲಿಸಿದ ಶ್ಲೋಕ ಮತ್ತು ವಿವರಣೆಗಳು."),
        ("auto_extracted", "Soft Blue (ನೀಲಿ)", "ಸ್ವಯಂಚಾಲಿತವಾಗಿ ಆಯ್ದುಕೊಂಡು ಕಾಲಂ ವಿಭಜನೆ ನಡೆಸಿದ ಪಠ್ಯ (ಚಿಕ್ಕಪುಟ್ಟ ಹೊಂದಾಣಿಕೆಗಳ ಕೊರತೆ ಇರಬಹುದು)."),
        ("phantom_disregard", "Soft Gray (ಬೂದು)", "ಫ್ಯಾಂಟಮ್ ಶ್ಲೋಕಗಳು — ಶ್ರೀ ಬನ್ನಂಜೆ ಗೋವಿಂದಾಚಾರ್ಯರ ಆವೃತ್ತಿಯಲ್ಲಿ ಈ ಶ್ಲೋಕ ಸಂಖ್ಯೆಗಳು ಅಸ್ತಿತ್ವದಲ್ಲಿಲ್ಲ. ಗೀತೆಯ ಸಾಂಪ್ರದಾಯಿಕ ಶ್ಲೋಕ ಸಂಖ್ಯೆಗಳನ್ನು ಕಾಯ್ದುಕೊಳ್ಳಲು ಮಾತ್ರ ಇವುಗಳನ್ನು ಖಾಲಿ ಇರಿಸಲಾಗಿದೆ.")
    ]
    
    bg_colors = {
        "clean": "FBF5E6",
        "screenshot_patch": "E2EFDA",
        "auto_extracted": "DDEBF7",
        "phantom_disregard": "F2F2F2"
    }
    
    for idx, (status, col_desc, desc) in enumerate(legend_data, start=1):
        c0 = legend_table.cell(idx, 0)
        c1 = legend_table.cell(idx, 1)
        c2 = legend_table.cell(idx, 2)
        
        c0.text = status
        c0.paragraphs[0].runs[0].font.bold = True
        
        c1.text = col_desc
        set_cell_background(c1, bg_colors[status])
        
        c2.text = desc
        
        for c in [c0, c1, c2]:
            set_cell_margins(c, top=80, bottom=80, left=120, right=120)
            
    doc.add_paragraph("\n")
    doc.add_page_break()

    # ---------------- 18 Chapters ----------------
    for ch in range(1, 19):
        ch_name = CHAPTER_NAMES[ch]
        ch_items = [x for x in data if x.get("chapter") == ch]
        
        print(f"Processing Chapter {ch}: {ch_name} ({len(ch_items)} verses)...")
        
        # Add Chapter Heading
        ch_p = doc.add_paragraph()
        ch_p.paragraph_format.space_before = Pt(18)
        ch_p.paragraph_format.space_after = Pt(12)
        ch_run = ch_p.add_run(f"ಅಧ್ಯಾಯ {ch} — {ch_name}")
        ch_run.font.name = 'Tiro Kannada'
        ch_run.font.size = Pt(22)
        ch_run.font.bold = True
        ch_run.font.color.rgb = RGBColor(0x8B, 0x1E, 0x1A)
        
        for item in ch_items:
            ref_val = item.get("ref", "")
            shloka_val = item.get("shloka", "").strip()
            meaning_val = item.get("meaning", "").strip()
            src_val = item.get("source_page", "")
            status_val = item.get("status", "")
            
            # Draw a beautiful Card as a single-cell table
            table = doc.add_table(rows=1, cols=1)
            table.autofit = True
            cell = table.cell(0, 0)
            
            # Padding
            set_cell_margins(cell, top=160, bottom=160, left=200, right=200)
            
            # Fill background based on status
            bg_color = bg_colors.get(status_val, "FFFFFF")
            set_cell_background(cell, bg_color)
            
            # 1. Header paragraph
            hp = cell.paragraphs[0]
            hp.paragraph_format.space_after = Pt(6)
            hrun = hp.add_run(f"ಶ್ಲೋಕ {ref_val}  |  Bannanje Book Page: {src_val}  |  [{status_val}]")
            hrun.font.bold = True
            hrun.font.size = Pt(9.5)
            hrun.font.color.rgb = RGBColor(0x7A, 0x6C, 0x54) # Muted beige-brown
            
            # 2. Sanskrit Shloka Block
            if shloka_val:
                sp = cell.add_paragraph()
                sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sp.paragraph_format.space_before = Pt(8)
                sp.paragraph_format.space_after = Pt(8)
                
                srun = sp.add_run(shloka_val)
                srun.font.name = 'Tiro Devanagari Sanskrit' if status_val == "clean" else 'Tiro Kannada'
                srun.font.size = Pt(11.5)
                srun.font.bold = True
                srun.font.color.rgb = RGBColor(0x1F, 0x1A, 0x13)
            
            # Divider line inside cell
            cell.add_paragraph("⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯⎯").alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[-1].runs[0].font.color.rgb = RGBColor(0xD6, 0xC5, 0x99)
            cell.paragraphs[-1].paragraph_format.space_after = Pt(6)
            
            # 3. Kannada Meaning & Commentary
            if meaning_val:
                # Handle paragraphs inside meaning
                paragraphs = meaning_val.split("\n\n")
                for p_text in paragraphs:
                    p_text = p_text.strip()
                    if not p_text:
                        continue
                    mp = cell.add_paragraph()
                    mp.paragraph_format.space_after = Pt(6)
                    mp.paragraph_format.line_spacing = 1.25
                    
                    mrun = mp.add_run(p_text)
                    mrun.font.name = 'Segoe UI'
                    mrun.font.size = Pt(10)
                    mrun.font.color.rgb = RGBColor(0x1F, 0x1A, 0x13)
                    
            # Space after table card
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(10)
            spacer.paragraph_format.space_after = Pt(10)
            
        doc.add_page_break()

    print(f"Saving Word document to {DOCX_PATH}...")
    doc.save(DOCX_PATH)
    print("Word document successfully rebuilt!")

if __name__ == "__main__":
    build_docx()
